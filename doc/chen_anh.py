"""Chen anh so do vao cac o placeholder trong Final Report .docx.

Tim o bang chua chuoi danh dau, xoa chu placeholder, chen anh canh giua voi be
rong vua khung noi dung trang A4. Anh qua cao thi rang buoc theo chieu cao de
khong tran sang trang sau.

    python doc/chen_anh.py doc/SIC_IoT_Final_Report_SmartHomeGarden.docx doc/hinh
"""
import copy
import struct
import sys
from pathlib import Path

import docx
import docx.table
import docx.text.paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Emu, Inches

DOCX = Path(sys.argv[1])
IMGDIR = Path(sys.argv[2])
# Doi so thu 3 (tuy chon): ghi ra file khac, dung khi ban goc dang mo trong Word.
OUT = Path(sys.argv[3]) if len(sys.argv) > 3 else DOCX

MAX_H = Inches(8.4)          # chua cho caption o cuoi trang

# chuoi nhan dang trong o placeholder -> ten file anh
MAP = [
    ('sơ đồ 6 (Gantt)',              '06-gantt-wbs.png'),
    ('sơ đồ 1 (Kiến trúc tổng quan)', '01-kien-truc-tong-quan.png'),
    ('sơ đồ 3 (Luồng xử lý dữ liệu)', '03-luong-xu-ly-du-lieu.png'),
    ('sơ đồ 2 (Cây quyết định 5 mức)', '02-cay-quyet-dinh-5-muc.png'),
    ('sơ đồ 5 (Triển khai)',          '05-so-do-trien-khai.png'),
    ('sơ đồ 7 trong',                 '07-so-do-chan-cam.png'),
    ('sơ đồ 4 (Sequence)',            '04a-sequence-tuoi-khan-cap.png'),
]

# khong dien: khong co du lieu that de chup / phai do nguoi chup
BO_QUA = [
    ('Chỗ dán ảnh chụp màn hình dashboard',
     'can anh dashboard luc CO du lieu that; anh trang rong se phan tac dung'),
    ('ẢNH NHÓM', 'anh chup nhom, chi nhom tu co'),
]


def png_size(p):
    d = p.read_bytes()
    return struct.unpack('>II', d[16:24])


def fit(img, content_w):
    """Tra (width, height) vua khung, uu tien lap day be rong."""
    w, h = png_size(img)
    width = content_w
    height = Emu(int(width * h / w))
    if height > MAX_H:
        height = MAX_H
        width = Emu(int(height * w / h))
    return width, height


def put_image(cell, img, content_w):
    """Xoa noi dung o roi chen anh canh giua."""
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    w, h = fit(img, content_w)
    p.add_run().add_picture(str(img), width=w, height=h)
    return w, h


def main():
    d = docx.Document(str(DOCX))
    sec = d.sections[0]
    content_w = Emu(sec.page_width - sec.left_margin - sec.right_margin)

    cells = [(t, c) for t in d.tables for r in t.rows for c in r.cells]
    done, missing = [], []
    tbl_seq = None          # giu lai bang hinh Figure 3.5 de nhan ban cho 3.6

    for marker, fname in MAP:
        img = IMGDIR / fname
        if not img.exists():
            missing.append(f'{fname}: khong co file')
            continue
        hit = next(((t, c) for t, c in cells if marker in c.text), None)
        if hit is None:
            missing.append(f'{fname}: khong thay o chua "{marker}"')
            continue
        tbl, cell = hit
        if marker == 'sơ đồ 4 (Sequence)':
            tbl_seq = tbl
        w, h = put_image(cell, img, content_w)
        done.append(f'{fname:<32} {Emu(w).inches:.2f} x {Emu(h).inches:.2f} in')

    # Them Figure moi cho so do 4b, vi bao cao chua co placeholder cho no.
    img4b = IMGDIR / '04b-sequence-watchdog.png'
    cap = next((p for p in d.paragraphs
                if p.text.strip().startswith('Figure 3.5')), None)
    if img4b.exists() and cap is not None and tbl_seq is not None \
            and 'Figure 3.6' not in d.element.xml:
        new_tbl = copy.deepcopy(tbl_seq._tbl)
        new_cap = copy.deepcopy(cap._p)
        cap._p.addnext(new_cap)
        cap._p.addnext(new_tbl)
        t2 = docx.table.Table(new_tbl, cap._parent)
        p2 = docx.text.paragraph.Paragraph(new_cap, cap._parent)
        w, h = put_image(t2.rows[0].cells[0], img4b, content_w)
        for r in list(p2.runs):
            r._element.getparent().remove(r._element)
        p2.add_run('Figure 3.6 — Serial watchdog releasing a supervisor-forced '
                   'actuator after the heartbeat stops.')
        done.append(f'{img4b.name:<32} {Emu(w).inches:.2f} x {Emu(h).inches:.2f} in  (Figure 3.6 moi)')

    d.save(str(OUT))

    print('=== DA CHEN ===')
    for x in done:
        print('  ' + x)
    if missing:
        print('\n=== KHONG CHEN DUOC ===')
        for x in missing:
            print('  ' + x)
    print('\n=== CO Y BO QUA ===')
    for marker, why in BO_QUA:
        still = any(marker in c.text for _, c in cells)
        print(f'  {marker[:42]:<44} {why}' + ('' if still else '  [!] khong con placeholder'))


if __name__ == '__main__':
    main()
